import re
from docx import Document

NEW_BASE_URL = "https://bob-eisenberg.com/Reprints/"

# Match only the old Rush URLs and capture everything after Reprints/
REPRINT_PATTERN = re.compile(
    r"https?://ftp\.rush\.edu/users/molebio/Bob_Eisenberg/Reprints/([^\s)\]]+)",
    re.IGNORECASE,
)


def replacement(match):
    """Construct the new URL using the captured filename/path."""
    return NEW_BASE_URL + match.group(1)


def update_paragraphs(paragraphs):
    """Replace visible URLs in a collection of paragraphs."""
    updates = 0

    for paragraph in paragraphs:
        for child in paragraph._element.iter():
            if child.tag.endswith("t") and child.text:
                new_text, count = REPRINT_PATTERN.subn(replacement, child.text)
                if count:
                    child.text = new_text
                    updates += count

    return updates


def update_tables(tables):
    """Recursively update URLs in all tables (including nested tables)."""
    updates = 0

    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                updates += update_paragraphs(cell.paragraphs)

                # Handle nested tables
                updates += update_tables(cell.tables)

    return updates


def update_cv_links(docx_path, output_path):
    doc = Document(docx_path)

    hyperlink_updates = 0
    text_updates = 0

    # ------------------------------------------------------------------
    # Step 1: Update hyperlink destinations (.rels)
    # ------------------------------------------------------------------
    for rel in doc.part.rels.values():
        if rel.reltype.endswith("/hyperlink"):
            match = REPRINT_PATTERN.search(rel.target_ref)
            if match:
                rel._target = replacement(match)
                hyperlink_updates += 1

    # ------------------------------------------------------------------
    # Step 2: Update visible URLs in the main document
    # ------------------------------------------------------------------
    text_updates += update_paragraphs(doc.paragraphs)
    text_updates += update_tables(doc.tables)

    # ------------------------------------------------------------------
    # Step 3: Update headers and footers
    # ------------------------------------------------------------------
    for section in doc.sections:
        text_updates += update_paragraphs(section.header.paragraphs)
        text_updates += update_tables(section.header.tables)

        text_updates += update_paragraphs(section.footer.paragraphs)
        text_updates += update_tables(section.footer.tables)

    # ------------------------------------------------------------------
    # Save the updated document
    # ------------------------------------------------------------------
    doc.save(output_path)

    print("\n--- Link Refactoring Metrics ---")
    print(f"✔ Modified hyperlink targets : {hyperlink_updates}")
    print(f"✔ Modified visible URLs      : {text_updates}")
    print(f"✔ Saved updated document to: {output_path}")


if __name__ == "__main__":
    update_cv_links(
        "../newest-CV/Bob_Eisenberg_CV_2026-02-19-2.docx",
        "../newest-CV/Bob_Eisenberg_CV_2026-02-19.docx",
    )
